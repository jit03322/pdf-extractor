from flask import Flask, request, jsonify,send_file
from flask_cors import CORS
import os
import pdfplumber
from docx import Document
import openpyxl
import re

app = Flask(__name__)
CORS(app)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Function to clean heading by keeping only alphabets and spaces
def clean_heading(heading):
    return re.sub(r'[^a-zA-Z\s]', '', heading).strip().lower()

# Function to extract content by detecting bold text as headings
def extract_content_by_boldness(pdf_path):
    content_list = []
    with pdfplumber.open(pdf_path) as pdf:
        current_heading = None
        current_content = []
        bold_words = []

        for page in pdf.pages:
            # Extract words with their font information
            words = page.extract_words(extra_attrs=["fontname", "size"])
            tables = page.extract_tables()

            for i, word in enumerate(words):
                # Check if the word is in bold (you may need to adapt the condition depending on the font)
                if "Bold" in word["fontname"]:
                    bold_words.append(word["text"])  
                    if (i + 1 < len(words) and "Bold" not in words[i + 1]["fontname"]) or i + 1 == len(words):
                        heading_text = " ".join(bold_words).strip()
                        bold_words = [] 

                        # Save the previous heading and content before capturing new heading
                        if current_heading:
                            
                            content_list.append((clean_heading(current_heading), current_content))
                        
                        current_heading = heading_text
                        current_content = []  
                else:
                    current_content.append(word["text"])

           
            if tables:
                for table in tables:
                    if current_heading:
                        content_list.append((clean_heading(current_heading), table))
                        current_content = []  
                    else:
                        content_list.append(('untitled_table', table))

        if current_heading and current_content:
            content_list.append((clean_heading(current_heading), current_content))

    return content_list


# Save paragraph to DOCX
def save_paragraph_to_docx(heading, content):
    doc = Document()
    doc.add_heading(heading, level=1)  
    doc.add_paragraph(" ".join(content))  
    doc.save(f'{heading}_output.docx')


# Save table to Excel
def save_table_to_excel(heading, table):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = heading

    # Populate the sheet with table data
    for row_index, row in enumerate(table, start=1):
        for col_index, cell in enumerate(row, start=1):
            sheet.cell(row=row_index, column=col_index).value = cell
    
    excel_path = f'{heading}_output.xlsx'
    workbook.save(excel_path)
    print(f"Table saved to {excel_path}")


# Upload route
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'pdf' not in request.files:
        return jsonify({'message': 'No file part'}), 400
    file = request.files['pdf']
    heading = request.form.get('heading') 
     # Get the user-provided heading
    
    if file.filename == '':
        return jsonify({'message': 'No selected file'}), 400

    pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(pdf_path)

    # Clean the user-provided heading for matching
    cleaned_user_heading = clean_heading(heading)
    print(f"Cleaned user-provided heading: {cleaned_user_heading}")

    # Extract content from the PDF based on bold headings
    extracted_content = extract_content_by_boldness(pdf_path)
    
    # Print all extracted headings and their content for debugging
    print("Extracted Headings and Content:")
    for extracted_heading, content in extracted_content:
        print(f"Heading: {extracted_heading}\nContent: {content}\n")

    # Check if cleaned user-provided heading matches any extracted cleaned headings
    for extracted_heading, content in extracted_content:
        if extracted_heading == cleaned_user_heading:  # Case-insensitive and cleaned match
            print(f"Match found for heading: {extracted_heading}")
            
            # Check if the content is a table or a paragraph
            if isinstance(content, list) and all(isinstance(row, list) for row in content):
    # This is a table (list of lists, where each list represents a row)
              print(f"Table found under heading '{extracted_heading}'. Saving to Excel.")
              save_table_to_excel(extracted_heading, content)
              return jsonify({'message': f'Table content saved to Excel for heading: {extracted_heading}.'})
            else:
    # If the content is not a table, assume it's a paragraph
              print(f"Paragraph found under heading '{extracted_heading}'. Saving to DOCX.")
              save_paragraph_to_docx(extracted_heading, content)
              return jsonify({'message': f'Paragraph content saved to DOCX for heading: {extracted_heading}.'})
    return jsonify({'message': 'No matching heading found in the PDF.'}), 404

# Start the server
if __name__ == '__main__':
    app.run(debug=True, port=5000) 